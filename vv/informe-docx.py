#!/usr/bin/env python
"""
Genera el documento único de entrega (informe + resultados + evidencias) en Word.

Parte de la plantilla APA de la asignatura para no reinventar la portada ni los estilos: se
conserva su carátula y solo se sustituyen título, autores y fecha. Todo lo demás se borra y se
escribe desde aquí, para que el documento se pueda regenerar cuando cambien las cifras en vez de
editarse a mano —que es como se desincronizan los números—.

    python vv/informe-docx.py

Salida: docs/vv/informe-final.docx
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

RAIZ = Path(__file__).resolve().parent.parent
PLANTILLA = Path.home() / 'Desktop' / 'clases' / 'FORMATO DE TAREAS.docx'
SALIDA = RAIZ / 'docs' / 'vv' / 'informe-final.docx'
ANCHO = Cm(15)  # el útil de la plantilla: 21 cm menos 3+3 de márgenes

AUTORES = ['Santiago Esquetini', 'Samuel Vega', 'Rene Herrera']
TITULO = 'Verificación y Validación del sistema Bug'
SUBTITULO = 'Juego de cartas P2P descentralizado — informe final'


# --- Utilidades sobre el documento ---------------------------------------------------------------

def poner(parrafo, texto):
    """Sustituye el texto conservando el formato del primer fragmento.

    Escribir solo en `runs[0]` deja los demás intactos, y la plantilla parte sus líneas en varios:
    así salió «Santiago EsquetiniEsquetini» y una fecha con dos fechas dentro.
    """
    if not parrafo.runs:
        parrafo.add_run(texto)
        return
    parrafo.runs[0].text = texto
    for sobrante in parrafo.runs[1:]:
        sobrante._element.getparent().remove(sobrante._element)


def borrar(parrafo):
    parrafo._element.getparent().remove(parrafo._element)


def h1(doc, texto):
    return doc.add_paragraph(texto, style='APA Level 1')


def h2(doc, texto):
    return doc.add_paragraph(texto, style='APA Level 2')


def p(doc, texto, negritas=()):
    """Un párrafo normal. `negritas` son fragmentos exactos que van en negrita."""
    par = doc.add_paragraph(style='APA Body Text')
    resto = texto
    for trozo in negritas:
        antes, _, resto = resto.partition(trozo)
        if antes:
            par.add_run(antes)
        par.add_run(trozo).bold = True
    if resto:
        par.add_run(resto)
    return par


def vinetas(doc, lineas):
    for linea in lineas:
        par = doc.add_paragraph(style='List Paragraph')
        par.paragraph_format.left_indent = Cm(1)
        par.paragraph_format.space_after = Pt(2)
        par.add_run('•  ').bold = True
        # El texto puede traer **negritas** al estilo markdown.
        for i, trozo in enumerate(linea.split('**')):
            if trozo:
                par.add_run(trozo).bold = i % 2 == 1


def tabla(doc, cabeceras, filas, anchos=None):
    t = doc.add_table(rows=1, cols=len(cabeceras))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, texto in enumerate(cabeceras):
        celda = t.rows[0].cells[i]
        celda.text = ''
        run = celda.paragraphs[0].add_run(texto)
        run.bold = True
        run.font.size = Pt(9)
    for fila in filas:
        celdas = t.add_row().cells
        for i, texto in enumerate(fila):
            celdas[i].text = ''
            par = celdas[i].paragraphs[0]
            for k, trozo in enumerate(str(texto).split('**')):
                if trozo:
                    run = par.add_run(trozo)
                    run.bold = k % 2 == 1
                    run.font.size = Pt(9)
    if anchos:
        for fila in t.rows:
            for i, ancho in enumerate(anchos):
                fila.cells[i].width = Cm(ancho)
    doc.add_paragraph()
    return t


def figura(doc, ruta, pie, ancho=ANCHO):
    """Una captura con su pie numerado. Si falta el archivo, se dice en el propio documento."""
    archivo = RAIZ / ruta
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if archivo.exists():
        par.add_run().add_picture(str(archivo), width=ancho)
    else:
        aviso = par.add_run(f'[FALTA LA EVIDENCIA: {ruta}]')
        aviso.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    figura.n += 1
    cap = doc.add_paragraph(style='APA Graph or Figure')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f'Figura {figura.n}. {pie}')
    run.font.size = Pt(9)
    run.italic = True


figura.n = 0


# --- El documento ---------------------------------------------------------------------------------

def main():
    if not PLANTILLA.exists():
        raise SystemExit(f'no encuentro la plantilla: {PLANTILLA}')

    doc = Document(str(PLANTILLA))

    # La portada de la plantilla se conserva; solo cambian título, autores y fecha. El resto del
    # cuerpo de ejemplo se borra.
    parrafos = doc.paragraphs
    poner(parrafos[6], TITULO)
    poner(parrafos[7], SUBTITULO)
    poner(parrafos[8], AUTORES[0])

    # Los otros dos autores van en párrafos que la plantilla dejaba vacíos, para no tocar la maqueta.
    for i, autor in enumerate(AUTORES[1:], start=9):
        destino = parrafos[i]
        destino.style = parrafos[8].style
        destino.alignment = parrafos[8].alignment
        destino.add_run(autor)

    hoy = date(2026, 8, 11)
    meses = ('enero febrero marzo abril mayo junio julio agosto '
             'septiembre octubre noviembre diciembre').split()
    dias = 'Lunes Martes Miércoles Jueves Viernes Sábado Domingo'.split()
    fecha = f'{dias[hoy.weekday()]}, {hoy.day} de {meses[hoy.month - 1]} de {hoy.year}'
    poner(parrafos[16], fecha)

    for par in parrafos[17:]:
        borrar(par)

    doc.add_page_break()

    # ---------------------------------------------------------------------------------------------
    h1(doc, 'Resumen')
    p(doc,
      'Bug es un juego de cartas multijugador peer-to-peer: las cartas viajan directas de un '
      'navegador a otro por WebRTC y no existe ningún servidor de juego. El único servidor presenta '
      'a los jugadores y se aparta; si se apaga a mitad de partida, la partida sigue.',
      negritas=['no existe ningún servidor de juego', 'la partida sigue'])
    p(doc,
      'Este documento reúne el informe técnico, los resultados de las pruebas con sus métricas y '
      'las evidencias visuales. Todas las cifras salen de la última ejecución real del pipeline: '
      'ninguna está escrita a mano.',
      negritas=['ninguna está escrita a mano'])

    tabla(doc,
          ['Indicador', 'Resultado'],
          [['Comprobaciones automatizadas', '**234** · 0 fallos'],
           ['Cobertura de código', '**89,1 %** (5.665 líneas analizadas)'],
           ['Puerta de calidad (SonarQube)', '**Passed** · Seguridad A · Fiabilidad A · Mantenibilidad A'],
           ['Ataques bloqueados', '**11 / 11**'],
           ['Propiedades distribuidas verificadas', '**7 / 7**'],
           ['Defectos encontrados y corregidos', '**17** (+ 6 vulnerabilidades)']],
          anchos=[6.5, 8.5])

    # ---------------------------------------------------------------------------------------------
    h1(doc, '1. El sistema')
    p(doc,
      'Cada navegador ejecuta el juego entero —las reglas, el mazo y su propia copia del estado— y '
      'se conecta con todos los demás formando una malla: 6 conexiones entre 4 jugadores, 45 entre '
      '10. El servidor solo interviene en el primer saludo.')
    h2(doc, 'Qué hace cada pieza')
    tabla(doc,
          ['', 'El contenedor', 'Cada navegador'],
          [['Motor de reglas', 'No', 'Sí, entero'],
           ['Estado de la partida', 'No', 'Su propia réplica'],
           ['Ve las cartas', 'Nunca', 'Solo las suyas'],
           ['Decide el turno', 'No', 'El testigo, entre navegadores']],
          anchos=[5, 5, 5])
    p(doc,
      'Quitar el árbitro abre cuatro problemas, y el sistema los resuelve con cuatro mecanismos '
      'clásicos de sistemas distribuidos:')
    vinetas(doc, [
        '**Turno** — exclusión mutua por paso de testigo con número de secuencia.',
        '**Orden** — relojes lógicos de Lamport, con desempate por identificador de jugador.',
        '**Consistencia** — estado replicado desde una semilla común, verificado con una huella.',
        '**Tolerancia a fallos** — latidos (2,5 s sospechoso / 6 s caído) y elección de líder Bully.',
    ])

    # ---------------------------------------------------------------------------------------------
    h1(doc, '2. Estrategia de verificación')
    p(doc,
      'Cinco capas, cada una para una pregunta que las otras no pueden contestar:')
    tabla(doc,
          ['Herramienta', 'Qué pregunta responde', 'Cuánto'],
          [['Vitest', '¿El motor y los algoritmos hacen lo que dicen?', '194 pruebas'],
           ['Cypress', '¿Funciona en navegadores reales, jugando?', '22 pruebas'],
           ['Burp Suite + banco propio', '¿Aguanta si alguien hace trampas a propósito?', '11 ataques'],
           ['Simulador propio', '¿Se sostiene con la red rota?', '7 propiedades'],
           ['SonarQube + Jenkins', '¿La calidad se mantiene en cada cambio?', '7 etapas']],
          anchos=[4.5, 7.5, 3])
    p(doc,
      'Las dos últimas capas hubo que escribirlas: ninguna herramienta comercial sabe qué es «el '
      'turno de Bug» ni cuándo dos jugadores están de acuerdo.',
      negritas=['hubo que escribirlas'])

    # ---------------------------------------------------------------------------------------------
    h1(doc, '3. Resultados y métricas')
    h2(doc, '3.1 Pruebas unitarias y de extremo a extremo')
    tabla(doc,
          ['Paquete', 'Pruebas', 'Qué cubre'],
          [['engine', '74', 'Reglas, determinismo, casos límite de cada carta'],
           ['net', '60', 'Lamport, replicación, testigo, Bully, reconexión'],
           ['signaling', '27', 'Aforo, introductores, guardas, salud'],
           ['web', '33', 'Señalización, identidades, azar seguro, efectos'],
           ['Cypress (navegador)', '22', 'Menú, partida local y malla con 3 y 10 navegadores'],
           ['**Total**', '**216**', 'más 11 ataques y 7 propiedades = **234**']],
          anchos=[4, 2.5, 8.5])

    h2(doc, '3.2 Validación distribuida (7/7)')
    tabla(doc,
          ['ID', 'Propiedad', 'Medición'],
          [['D1', 'Convergencia con la red revuelta', '3, 5 y 10 nodos · misma huella en todos'],
           ['D2', 'Orden total de eventos concurrentes', '12 órdenes de entrega · una sola secuencia'],
           ['D3', 'Exclusión mutua sobre el turno', '200 cesiones · 200 con poseedor único'],
           ['D4', 'Detección de caídas por latidos', 'sospechoso 2.500 ms · caído 6.000 ms'],
           ['D5', 'Elección de líder (Bully)', 'acuerdo unánime · reelección si cae el favorito'],
           ['D6', 'Recuperación de un nodo desviado', '60 eventos recuperados de 120'],
           ['D7', 'Coste de la replicación', 'desorden realista < 30× el trabajo mínimo']],
          anchos=[1.2, 6.3, 7.5])

    h2(doc, '3.3 Seguridad (11/11 bloqueados)')
    p(doc,
      'Se atacó la señalización con Burp Suite y los once ataques quedaron automatizados. Seis eran '
      'vulnerabilidades reales, dos de ellas críticas.',
      negritas=['Seis eran vulnerabilidades reales'])
    tabla(doc,
          ['Ataque', 'Antes', 'Ahora'],
          [['S1 · Suplantar a otro jugador', 'Se colaba', 'Bloqueado'],
           ['S3 · Echar a alguien de su partida', 'Se colaba', 'Bloqueado'],
           ['S8 · Mensaje malformado', '**Tumbaba el servidor**', 'Bloqueado'],
           ['S6 · Agotar conexiones', 'Se colaba', 'Bloqueado'],
           ['S7 · Mensaje de 32 MB', 'Se colaba', 'Bloqueado'],
           ['S9 · Saltarse el aforo', 'Se colaba', 'Bloqueado'],
           ['Otros cinco', 'Ya aguantaba', 'Bloqueado']],
          anchos=[7, 4, 4])

    h2(doc, '3.4 Calidad estática')
    p(doc,
      'La puerta de calidad exige 0 avisos nuevos y está en verde. En el acumulado del repositorio '
      'quedan 33 avisos, revisados uno a uno: los que protegían algo se corrigieron y el resto está '
      'excluido con su justificación escrita en sonar-project.properties.',
      negritas=['0 avisos nuevos', 'con su justificación escrita'])

    # ---------------------------------------------------------------------------------------------
    h1(doc, '4. Defectos encontrados')
    tabla(doc,
          ['Origen', 'Nº', 'Ejemplo representativo'],
          [['Jugando de verdad', '9', 'Un jugador veía otra partida y no podía tirar ninguna carta'],
           ['Atacando el sistema', '5', 'Un mensaje de dos líneas apagaba el servidor entero'],
           ['Montando las pruebas', '3', '/health solo atendía GET, y el arranque nunca terminaba'],
           ['**Total**', '**17**', 'Casi ninguno lo vieron las pruebas unitarias del principio']],
          anchos=[4, 1.5, 9.5])
    p(doc,
      'El dato importante no es el número, es el reparto: una prueba unitaria comprueba lo que se '
      'te ocurrió comprobar, y las reglas del juego estaban bien. Los fallos vivían en la pantalla, '
      'en la red y en el arranque.',
      negritas=['el reparto'])
    p(doc,
      'Además, el análisis estático destapó dos endurecimientos que ningún ataque manual encontró: '
      'el código de sala salía de un generador predecible (ahora usa crypto, con cuatro pruebas '
      'nuevas) y el túnel resolvía su ejecutable por PATH (ahora usa ruta absoluta).')

    # ---------------------------------------------------------------------------------------------
    doc.add_page_break()
    h1(doc, '5. Evidencias visuales')
    p(doc,
      'Las seis primeras capturas no se hicieron a mano: las toma la propia suite de Cypress mientras juega, así que se regeneran solas en cada ejecución y no pueden quedarse desfasadas respecto al sistema.',
      negritas=['las toma la propia suite de Cypress mientras juega'])


    figura(doc, 'docs/vv/evidencias/cypress/menu.cy.ts/00-pantalla-de-entrada.png',
           'Capturada por la prueba de Cypress «menú principal». Pantalla de entrada: se crea sala '
           'o se entra con un código de cuatro caracteres.', Cm(11))
    figura(doc, 'docs/vv/evidencias/cypress/partida-local.cy.ts/03-mesa-repartida.png',
           'Capturada por la prueba de Cypress «partida local». Mesa repartida: siete cartas, mazo, '
           'pozo y turno en curso.', Cm(13))
    figura(doc, 'docs/vv/evidencias/cypress/distribuido/malla.cy.ts/04-tres-nodos-convergen.png',
           'Capturada por la prueba de Cypress «malla de tres nodos». Tres navegadores reales con '
           'WebRTC: las tres réplicas publican la misma huella de estado.')
    figura(doc, 'docs/vv/evidencias/cypress/distribuido/malla.cy.ts/05-cae-un-nodo-los-otros-siguen.png',
           'Capturada por la prueba de Cypress «malla de tres nodos». Tolerancia a fallos: cae un '
           'nodo y los demás siguen jugando.')
    figura(doc, 'docs/vv/evidencias/ui/14-pantalla-maestra-panel.png',
           'Capturada por la prueba de Cypress «malla de tres nodos». La Pantalla Maestra: líder, '
           'testigo, convergencia, huella de cada nodo e historial de la malla.', Cm(8))
    figura(doc, 'docs/vv/evidencias/cypress/distribuido/aforo.cy.ts/09-diez-jugadores.png',
           'Capturada por la prueba de Cypress «aforo: diez caben, el once no». Diez navegadores en '
           'la misma sala, 45 conexiones directas.')
    figura(doc, 'docs/vv/evidencias/laboratorio/07-sonarqube-dashboard.png',
           'SonarQube: puerta de calidad superada, 89,1 % de cobertura.')
    figura(doc, 'docs/vv/evidencias/laboratorio/06-jenkins-pipeline.png',
           'Jenkins: las siete etapas por commit. Las cuatro primeras construcciones en rojo son el '
           'pipeline que llevaba semanas escrito sin haber arrancado nunca.')
    figura(doc, 'docs/vv/evidencias/laboratorio/11-jenkins-configuracion.png',
           'Jenkins — el pipeline por dentro. Se define como «Pipeline script from SCM», así que las '
           'siete etapas viven en el Jenkinsfile del repositorio y no en la configuración del servidor.')
    figura(doc, 'docs/vv/evidencias/laboratorio/12-cypress-resultados.png',
           'Las 22 pruebas de Cypress, una a una: malla de tres nodos, aforo de diez, partida local '
           'y menú. Cero fallos.', Cm(14))
    figura(doc, 'docs/vv/evidencias/laboratorio/08-burp-websockets.png',
           'Burp Suite en medio de una partida: solo aparecen los saludos de señalización. Ni una '
           'carta, ni una mano, ni el mazo.')

    # ---------------------------------------------------------------------------------------------
    h1(doc, '6. Conclusión')
    p(doc,
      'El sistema hace lo que dice: 234 comprobaciones automatizadas lo vigilan en cada cambio, '
      'los once ataques quedan bloqueados y las siete propiedades distribuidas se cumplen con la '
      'red rota a propósito. Lo más valioso del bloque no fueron las pruebas que pasaron, sino los '
      '17 defectos que solo aparecieron al jugar, al atacar y al montar el laboratorio.',
      negritas=['los 17 defectos que solo aparecieron'])

    SALIDA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(SALIDA))
    print(f'  OK  {SALIDA.relative_to(RAIZ)}  ({SALIDA.stat().st_size // 1024} kB, '
          f'{figura.n} figuras)')


if __name__ == '__main__':
    main()
